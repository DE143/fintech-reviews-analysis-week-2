# create_professional_report.py
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from datetime import datetime
import os

def create_professional_report():
    """Create a professional .docx report combining all files - FIXED VERSION"""
    
    print("🎨 Creating professional DOCX report...")
    
    # Create document
    doc = Document()
    
    # ==================== TITLE PAGE ====================
    print("📝 Creating title page...")
    
    # Title
    title = doc.add_heading('CUSTOMER EXPERIENCE ANALYTICS\nFOR FINTECH APPLICATIONS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for run in title.runs:
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(0, 51, 102)
        run.bold = True
    
    # Subtitle
    subtitle = doc.add_heading('Comprehensive Analysis Report', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for run in subtitle.runs:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0, 102, 204)
        run.italic = True
    
    doc.add_paragraph()
    
    # Company
    company = doc.add_paragraph('Derese Ewunet')
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for run in company.runs:
        run.font.size = Pt(18)
        run.bold = True
    
    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f'Report Date: {datetime.now().strftime("%B %d, %Y")}')
    date_run.font.size = Pt(12)
    date_run.italic = True
    
    doc.add_page_break()
    
    # ==================== TABLE OF CONTENTS ====================
    print("📑 Creating table of contents...")
    
    toc_title = doc.add_heading('TABLE OF CONTENTS', 1)
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # TOC Items - SIMPLIFIED
    doc.add_paragraph()
    
    toc_items = [
        "1. EXECUTIVE SUMMARY",
        "2. PROJECT OVERVIEW",
        "3. METHODOLOGY",
        "4. DATA ANALYSIS RESULTS",
        "   4.1 Performance Metrics",
        "   4.2 Sentiment Analysis",
        "   4.3 Thematic Insights",
        "5. VISUAL ANALYSIS",
        "   5.1 Rating Distribution",
        "   5.2 Sentiment Comparison",
        "   5.3 Performance Benchmarking",
        "6. BANK-SPECIFIC ANALYSIS",
        "   6.1 Commercial Bank of Ethiopia",
        "   6.2 Bank of Abyssinia",
        "   6.3 Dashen Bank",
        "7. ACTIONABLE RECOMMENDATIONS",
        "8. EXPECTED BUSINESS IMPACT",
        "9. TECHNICAL IMPLEMENTATION",
        "10. CONCLUSION",
        "APPENDICES",
        "   A. Data Summary",
        "   B. Detailed Statistics"
    ]
    
    for item in toc_items:
        if item.startswith(' '):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.5)
            run = para.add_run(item[4:])  # Remove 4 spaces for indentation
        else:
            para = doc.add_paragraph(item)
            for run in para.runs:
                run.bold = True
    
    doc.add_page_break()
    
    # ==================== EXECUTIVE SUMMARY ====================
    print("📊 Creating executive summary...")
    
    section_title = doc.add_heading('1. EXECUTIVE SUMMARY', 1)
    
    # Key Findings Table - FIXED: Using only 4 data rows + header = 5 total rows
    findings_data = [
        ('Performance Leader', 'Dashen Bank (4.04/5.0)'),
        ('Sentiment Leader', 'Dashen Bank (68.2% Positive)'),
        ('Improvement Priority', 'Bank of Abyssinia (2.11/5.0)'),
        ('Data Coverage', '1,500+ Reviews Analyzed')
    ]
    
    # Create table with correct number of rows: header + data rows
    findings_table = doc.add_table(rows=len(findings_data) + 1, cols=2)
    findings_table.style = 'Light Grid Accent 1'
    findings_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    header_cells = findings_table.rows[0].cells
    header_cells[0].text = 'KEY METRIC'
    header_cells[1].text = 'FINDING'
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Data rows - CORRECT: i starts from 1, not 0
    for i, (metric, finding) in enumerate(findings_data, 1):  # Start from 1 for data rows
        row_cells = findings_table.rows[i].cells  # This is now safe
        row_cells[0].text = metric
        row_cells[1].text = finding
    
    doc.add_paragraph()
    
    # Summary text
    summary_text = """
    This comprehensive analysis provides data-driven insights into customer experience across three major Ethiopian banking applications. 
    The study examines 1,500+ Google Play Store reviews using advanced Natural Language Processing (NLP) techniques and statistical analysis.
    
    Key insights reveal significant performance variations, with Dashen Bank demonstrating exceptional user satisfaction while 
    Commercial Bank of Ethiopia and Bank of Abyssinia require targeted improvements in performance, functionality, and user experience.
    
    The report provides actionable recommendations for each banking institution, supported by visual evidence and statistical validation.
    """
    
    for paragraph in summary_text.strip().split('\n\n'):
        doc.add_paragraph(paragraph)
    
    doc.add_page_break()
    
    # ==================== PROJECT OVERVIEW ====================
    print("📋 Creating project overview...")
    
    section_title = doc.add_heading('2. PROJECT OVERVIEW', 1)
    
    overview_text = """
    2.1 OBJECTIVES
    
    • Identify key drivers of customer satisfaction and dissatisfaction
    • Provide data-driven product improvement recommendations
    • Enable competitive benchmarking across banking applications
    • Support strategic roadmap development for digital banking
    
    2.2 BANKS ANALYZED
    
    • Commercial Bank of Ethiopia (CBE): Market leader with 5M+ installs
    • Bank of Abyssinia (BOA): Growing digital banking presence
    • Dashen Bank: Technology-forward approach with modern interface
    
    2.3 SCOPE
    
    The analysis covers comprehensive aspects including:
    • Performance metrics and rating distributions
    • Sentiment analysis using advanced NLP techniques
    • Thematic analysis of user feedback
    • Comparative benchmarking across applications
    • Actionable improvement recommendations
    """
    
    for paragraph in overview_text.strip().split('\n\n'):
        if paragraph.startswith('2.'):
            doc.add_heading(paragraph.split(' ')[1], level=2)
            content = paragraph.split(' ', 2)[2] if len(paragraph.split(' ', 2)) > 2 else ""
            if content:
                doc.add_paragraph(content)
        elif paragraph.startswith('•'):
            for line in paragraph.split('\n'):
                if line.strip():
                    doc.add_paragraph(line, style='List Bullet')
        else:
            doc.add_paragraph(paragraph)
    
    doc.add_page_break()
    
    # ==================== METHODOLOGY ====================
    print("🔧 Adding methodology...")
    
    doc.add_heading('3. METHODOLOGY', 1)
    
    methodology_text = """
    3.1 DATA COLLECTION
    
    • Source: Google Play Store official reviews
    • Sample Size: 1,500+ reviews (500 per bank)
    • Collection Period: Recent user reviews
    • Tools: google-play-scraper Python library
    
    3.2 ANALYTICAL APPROACH
    
    • Sentiment Analysis: DistilBERT transformer model
    • Thematic Analysis: TF-IDF keyword extraction
    • Statistical Analysis: Rating distributions and correlations
    • Database: PostgreSQL for structured storage
    
    3.3 QUALITY ASSURANCE
    
    • Data validation at each processing stage
    • Cross-validation of analytical results
    • Statistical significance testing
    • Reproducibility verification
    """
    
    for paragraph in methodology_text.strip().split('\n\n'):
        if paragraph.startswith('3.'):
            doc.add_heading(paragraph.split(' ')[1], level=2)
            content = paragraph.split(' ', 2)[2] if len(paragraph.split(' ', 2)) > 2 else ""
            if content:
                doc.add_paragraph(content)
        elif paragraph.startswith('•'):
            for line in paragraph.split('\n'):
                if line.strip():
                    doc.add_paragraph(line, style='List Bullet')
    
    doc.add_page_break()
    
    # ==================== VISUAL ANALYSIS ====================
    print("📈 Adding visualizations...")
    
    doc.add_heading('5. VISUAL ANALYSIS', 1)
    
    # Add rating distribution chart
    if os.path.exists('reports/rating_distribution.png'):
        doc.add_heading('5.1 Rating Distribution Analysis', 2)
        doc.add_picture('reports/rating_distribution.png', width=Inches(6))
        caption = doc.add_paragraph('Figure 1: Rating distribution across banking applications')
        caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs:
            run.italic = True
    
    doc.add_paragraph()
    
    # Add sentiment analysis chart
    if os.path.exists('reports/sentiment_analysis.png'):
        doc.add_heading('5.2 Sentiment Analysis Comparison', 2)
        doc.add_picture('reports/sentiment_analysis.png', width=Inches(6))
        caption = doc.add_paragraph('Figure 2: Comparative sentiment analysis across banks')
        caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs:
            run.italic = True
    
    doc.add_paragraph()
    
    # Add average rating chart
    if os.path.exists('reports/average_rating.png'):
        doc.add_heading('5.3 Performance Benchmarking', 2)
        doc.add_picture('reports/average_rating.png', width=Inches(6))
        caption = doc.add_paragraph('Figure 3: Average rating comparison with performance thresholds')
        caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs:
            run.italic = True
    
    doc.add_page_break()
    
    # ==================== BANK ANALYSIS ====================
    print("🏦 Adding bank analysis...")
    
    doc.add_heading('6. BANK-SPECIFIC ANALYSIS', 1)
    
    # Create performance comparison table
    perf_table = doc.add_table(rows=4, cols=4)
    perf_table.style = 'Light Grid Accent 2'
    perf_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header
    headers = ['BANK', 'AVG RATING', 'POSITIVE %', 'PERFORMANCE']
    header_cells = perf_table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for run in header_cells[i].paragraphs[0].runs:
            run.bold = True
    
    # Data rows - CORRECT: Only 3 data rows
    banks_data = [
        ('Commercial Bank of Ethiopia', '2.70/5.0', '19.8%', 'Needs Improvement'),
        ('Bank of Abyssinia', '2.11/5.0', '18.4%', 'Priority Improvement'),
        ('Dashen Bank', '4.04/5.0', '68.2%', 'Excellent')
    ]
    
    for i, (bank, rating, positive, performance) in enumerate(banks_data, 1):  # Start from 1
        row_cells = perf_table.rows[i].cells
        row_cells[0].text = bank
        row_cells[1].text = rating
        row_cells[2].text = positive
        row_cells[3].text = performance
    
    doc.add_paragraph()
    
    # Bank-specific insights
    bank_insights = [
        ("6.1 Commercial Bank of Ethiopia", [
            "Strengths: Comprehensive feature set, strong brand recognition",
            "Challenges: Performance optimization, transaction speed",
            "Opportunities: Error handling improvements, UX simplification"
        ]),
        ("6.2 Bank of Abyssinia", [
            "Strengths: Digital innovation focus, modern technology stack",
            "Challenges: Critical performance issues, poor user experience",
            "Opportunities: Complete UI redesign, authentication improvements"
        ]),
        ("6.3 Dashen Bank", [
            "Strengths: Excellent UI design, high customer satisfaction",
            "Challenges: Limited advanced features, minor performance issues",
            "Opportunities: Feature expansion, support enhancement"
        ])
    ]
    
    for title, insights in bank_insights:
        doc.add_heading(title, 2)
        for insight in insights:
            doc.add_paragraph(f"• {insight}", style='List Bullet')
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ==================== RECOMMENDATIONS ====================
    print("🎯 Adding recommendations...")
    
    doc.add_heading('7. ACTIONABLE RECOMMENDATIONS', 1)
    
    recommendations = [
        ("IMMEDIATE PRIORITIES (0-3 MONTHS)", [
            "Commercial Bank of Ethiopia: Optimize transaction processing",
            "Bank of Abyssinia: Complete user interface redesign",
            "Dashen Bank: Expand advanced feature set"
        ]),
        ("STRATEGIC INITIATIVES (3-12 MONTHS)", [
            "Implement AI-powered financial insights",
            "Develop personalized banking dashboards",
            "Integrate proactive customer support"
        ])
    ]
    
    for title, recs in recommendations:
        doc.add_heading(title, 2)
        for rec in recs:
            doc.add_paragraph(f"• {rec}", style='List Bullet')
        doc.add_paragraph()
    
    # ==================== BUSINESS IMPACT ====================
    print("📈 Adding business impact...")
    
    doc.add_heading('8. EXPECTED BUSINESS IMPACT', 1)
    
    impact_points = [
        "20-30% reduction in negative reviews within 6 months",
        "15-25% improvement in app store ratings",
        "10-20% increase in daily active users",
        "Enhanced customer retention and satisfaction",
        "Competitive advantage through data-driven decisions"
    ]
    
    for point in impact_points:
        doc.add_paragraph(f"• {point}", style='List Bullet')
    
    doc.add_page_break()
    
    # ==================== DATA APPENDIX ====================
    print("📊 Adding data appendix...")
    
    doc.add_heading('APPENDICES', 1)
    
    # Try to load and include the actual text files
    appendix_files = [
        ('A. EXECUTIVE SUMMARY', 'reports/executive_summary.txt'),
        ('B. DATA SUMMARY', 'reports/data_summary.txt'),
        ('C. RECOMMENDATIONS', 'reports/recommendations.txt')
    ]
    
    for title, filepath in appendix_files:
        doc.add_heading(title, 2)
        
        if os.path.exists(filepath):
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    content = f.read()
                # Add content with simple formatting
                for line in content.split('\n'):
                    if line.strip():
                        if line.startswith('=') or line.startswith('-'):
                            # Skip decorative lines
                            continue
                        elif line.isupper() and len(line) < 50:
                            # Likely a heading
                            para = doc.add_paragraph(line)
                            for run in para.runs:
                                run.bold = True
                        else:
                            doc.add_paragraph(line)
                doc.add_paragraph()
            except:
                doc.add_paragraph(f"[Content from {filepath} could not be loaded]")
        else:
            doc.add_paragraph(f"[File not found: {filepath}]")
        
        doc.add_paragraph()
    
    # ==================== CONCLUSION ====================
    print("🎓 Adding conclusion...")
    
    doc.add_heading('10. CONCLUSION', 1)
    
    conclusion_text = """
    This comprehensive analysis provides a clear roadmap for mobile banking application improvements across Ethiopian financial institutions. 
    By leveraging data-driven insights, each bank can prioritize enhancements that will deliver maximum impact on customer satisfaction.
    
    The digital transformation of Ethiopia's banking sector presents significant opportunities for innovation and customer experience enhancement. 
    This report establishes a foundation for evidence-based decision-making and strategic improvement planning.
    """
    
    for paragraph in conclusion_text.strip().split('\n\n'):
        doc.add_paragraph(paragraph)
    
    doc.add_paragraph()
    
    next_steps = doc.add_heading('NEXT STEPS', 2)
    
    steps = [
        "1. Present findings to bank stakeholders",
        "2. Develop detailed implementation roadmaps", 
        "3. Establish monitoring and measurement frameworks",
        "4. Schedule follow-up analysis in 6 months"
    ]
    
    for step in steps:
        doc.add_paragraph(step)
    
    # ==================== FINAL PAGE ====================
    print("📝 Adding final page...")
    
    doc.add_page_break()
    
    final_title = doc.add_heading('Derese Ewunet', 0)
    final_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contact_info = """
    For more information or to discuss implementation:
    
    Email: derese641735.ew@gmail.com
    Phone: +251 943 482 726
    Address: Addis Ababa, Ethiopia
    
    CONFIDENTIAL - INTERNAL USE ONLY
    
    Report Generated: """ + datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    contact_para = doc.add_paragraph(contact_info)
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ==================== SAVE DOCUMENT ====================
    output_path = 'reports/Customer_Experience_Analytics_Final_Report_Professional.docx'
    doc.save(output_path)
    
    print(f"\n✅ Professional DOCX report saved: {output_path}")
    print("="*60)
    print("\n📋 REPORT CONTENTS:")
    print("  • Title Page with Professional Branding")
    print("  • Table of Contents")
    print("  • Executive Summary with Key Findings")
    print("  • Project Overview & Methodology")
    print("  • Visual Analysis with Embedded Charts")
    print("  • Bank-Specific Performance Analysis")
    print("  • Actionable Recommendations")
    print("  • Expected Business Impact")
    print("  • Data Appendices")
    print("  • Conclusion & Next Steps")
    print("\n🎯 The report is now ready for professional submission!")

if __name__ == "__main__":
    try:
        create_professional_report()
    except ImportError:
        print("\n❌ ERROR: python-docx is not installed!")
        print("Please install it using: pip install python-docx")
    except Exception as e:
        print(f"\n❌ ERROR: {e}")
        import traceback
        traceback.print_exc()